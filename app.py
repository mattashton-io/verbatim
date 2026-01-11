import os
import uuid
import threading
import time
import io
import re
from flask import Flask, request, render_template, jsonify, send_from_directory, send_file
from google.cloud import storage
from google.cloud import secretmanager
from transcription_service_v1 import transcribe_gcs_file, refine_text_with_gemini
from dotenv import load_dotenv
from docx import Document

# Load environment variables from .env if it exists
load_dotenv()

app = Flask(__name__)

# Configuration
GCS_BUCKET_NAME = os.environ.get('GCS_BUCKET_NAME')
PROJECT_ID = os.environ.get('GOOGLE_CLOUD_PROJECT')


# Simple in-memory job store (for demo purposes)
jobs = {}

def upload_to_gcs(file_obj, filename, bucket_name, content_type=None):
    """Uploads a file object to Google Cloud Storage and returns the gs:// URI."""
    print(f"[DEBUG] Starting upload to GCS bucket '{bucket_name}' for file: {filename}")
    start_time = time.time()
    
    client = storage.Client()
    bucket = client.bucket(bucket_name)
    blob_name = f"uploads/{uuid.uuid4()}-{filename}"
    blob = bucket.blob(blob_name)
    
    if content_type:
        blob.upload_from_file(file_obj, content_type=content_type)
    else:
        blob.upload_from_file(file_obj)
    
    elapsed = time.time() - start_time
    print(f"[DEBUG] Upload to GCS completed in {elapsed:.2f}s. URI: gs://{bucket_name}/{blob_name}")
    
    return f"gs://{bucket_name}/{blob_name}"

def run_transcription_pipeline(job_id, gcs_uri):
    """Background task to process audio transcription."""
    print(f"[DEBUG] Job {job_id}: Starting transcription pipeline for {gcs_uri}")
    try:
        # Stage 1: Transcription
        jobs[job_id]['status'] = 'transcribing'
        print(f"[DEBUG] Job {job_id}: Calling transcribe_gcs_file (V1 Async)...")
        
        transcribe_start = time.time()
        raw_text = transcribe_gcs_file(gcs_uri)
        transcribe_duration = time.time() - transcribe_start
        
        print(f"[DEBUG] Job {job_id}: Transcription complete in {transcribe_duration:.2f}s. Raw text length: {len(raw_text)} chars")
        
        # Stage 2: Refinement with Gemini 3
        jobs[job_id]['status'] = 'refining'
        print(f"[DEBUG] Job {job_id}: Calling refine_text_with_gemini (Gemini 3 Flash)...")
        
        refine_start = time.time()
        refined_text = refine_text_with_gemini(raw_text)
        refine_duration = time.time() - refine_start
        
        print(f"[DEBUG] Job {job_id}: Refining complete in {refine_duration:.2f}s. Final text length: {len(refined_text)} chars")
        
        # Success
        jobs[job_id]['status'] = 'completed'
        jobs[job_id]['result'] = refined_text
        print(f"[DEBUG] Job {job_id}: Pipeline finished successfully.")
        
    except Exception as e:
        print(f"[ERROR] Job {job_id}: Error in transcription pipeline: {e}")
        jobs[job_id]['status'] = 'failed'
        jobs[job_id]['error'] = str(e)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload():
    print("[DEBUG] Received /upload request")
    if 'audio_file' not in request.files:
        print("[DEBUG] No 'audio_file' in request.files")
        return jsonify({"error": "No file part"}), 400
    
    file = request.files['audio_file']
    if file.filename == '':
        print("[DEBUG] Empty filename")
        return jsonify({"error": "No selected file"}), 400
    
    if not GCS_BUCKET_NAME:
        print("[ERROR] GCS_BUCKET_NAME env var not set")
        return jsonify({"error": "GCS_BUCKET_NAME not configured"}), 500

    try:
        filename = file.filename
        
        # Explicitly reject M4A due to Python 3.13 / audioop limitations
        if filename.lower().endswith('.m4a'):
            print(f"[DEBUG] Rejected M4A file: {filename}")
            return jsonify({"error": "M4A format is not currently supported in Python 3.13. Please upload .mp3 or .wav"}), 400

        print(f"[DEBUG] Processing file for upload: {filename}")

        # Step 1: Upload to GCS
        gcs_uri = upload_to_gcs(file, filename, GCS_BUCKET_NAME)
        
        # Step 2: Create a Job ID for polling
        job_id = str(uuid.uuid4())
        jobs[job_id] = {'status': 'pending', 'uri': gcs_uri}
        print(f"[DEBUG] Job created: {job_id}")
        
        # Step 3: Trigger background processing in a separate thread
        thread = threading.Thread(target=run_transcription_pipeline, args=(job_id, gcs_uri))
        thread.start()
        print(f"[DEBUG] Background thread started for job {job_id}")
        
        return jsonify({"job_id": job_id, "gcs_uri": gcs_uri})
    except Exception as e:
        print(f"[ERROR] Upload/Processing error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/status/<job_id>')
def status(job_id):
    # Log polling occasionally if needed for debugging
    # print(f"[DEBUG] Status check for job {job_id}") 
    job = jobs.get(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404
    return jsonify(job)

def create_docx(text):
    doc = Document()
    doc.add_heading('Transcript', 0)
    
    # Split by <br> tags (handling reasonable variations)
    paragraphs = re.split(r'<br\s*/?>', text)
    
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para.strip())
            
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

@app.route('/download/<job_id>/<fmt>')
def download_transcript(job_id, fmt):
    job = jobs.get(job_id)
    if not job or job.get('status') != 'completed':
        return jsonify({"error": "Job not completed or found"}), 404
    
    refined_text = job.get('result', '')
    
    if fmt == 'txt':
        # Simple text, replace <br> with newlines
        text_content = refined_text.replace("<br>", "\n").replace("<br/>", "\n")
        bio = io.BytesIO(text_content.encode('utf-8'))
        return send_file(
            bio,
            as_attachment=True,
            download_name=f"transcript.txt", # simplified name
            mimetype='text/plain'
        )
    elif fmt == 'md':
        # Markdown, replace <br> with \n\n
        md_content = f"# Transcript\n\n{refined_text.replace('<br>', '\n\n').replace('<br/>', '\n\n')}"
        bio = io.BytesIO(md_content.encode('utf-8'))
        return send_file(
            bio,
            as_attachment=True,
            download_name=f"transcript.md",
            mimetype='text/markdown'
        )
    elif fmt == 'docx':
        try:
            bio = create_docx(refined_text)
            return send_file(
                bio,
                as_attachment=True,
                download_name=f"transcript.docx",
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        except Exception as e:
            print(f"Error creating DOCX: {e}")
            return jsonify({"error": str(e)}), 500
    else:
        return jsonify({"error": "Invalid format"}), 400

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8080))
    print(f"[DEBUG] Starting Verbatim app on port {port}...")
    app.run(debug=True, host='0.0.0.0', port=port)